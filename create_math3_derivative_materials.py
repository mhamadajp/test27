from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt


OUT_DIR = Path("数学III_第03回_微分係数教材")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def set_row_height(row, height_cm):
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = OxmlElement("w:trHeight")
    tr_h.set(qn("w:val"), str(int(height_cm * 567)))
    tr_h.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_h)


def setup_doc(title):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Yu Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Yu Gothic"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    return doc


def add_footer(doc, label):
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.text = label
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.add_run(title)


def add_subsection(doc, title):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.add_run(title)


def add_note_box(doc, title, lines, fill="EEF6FF"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    for line in lines:
        p = cell.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def add_write_box(doc, prompts, height_cm=2.2):
    table = doc.add_table(rows=len(prompts), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, prompt in enumerate(prompts):
        row = table.rows[i]
        set_row_height(row, height_cm)
        row.cells[0].text = prompt
        row.cells[1].text = ""
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_width(row.cells[0], 5.2)
        set_cell_width(row.cells[1], 12.0)
        set_cell_shading(row.cells[0], "F8FAFC")
    doc.add_paragraph()


def add_checklist(doc, items):
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    for row, item in zip(table.rows, items):
        row.cells[0].text = "□"
        row.cells[1].text = item
        set_cell_width(row.cells[0], 1.0)
        set_cell_width(row.cells[1], 16.0)
    doc.add_paragraph()


def m_text(text):
    return f"<m:r><m:t>{escape(text)}</m:t></m:r>"


def m_frac(num, den):
    return f"<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>"


def m_sup(base, sup):
    return f"<m:sSup><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>"


def m_sub(base, sub):
    return f"<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>"


def m_sqrt(content):
    return f'<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr><m:e>{content}</m:e></m:rad>'


def m_lim_low(limit_text):
    return f"<m:limLow><m:e>{m_text('lim')}</m:e><m:lim>{m_text(limit_text)}</m:lim></m:limLow>"


def add_equation(doc, parts, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    xml = f'<m:oMath {nsdecls("m")}>{"".join(parts)}</m:oMath>'
    p._p.append(parse_xml(xml))
    return p


def add_inline_equation(paragraph, parts):
    xml = f'<m:oMath {nsdecls("m")}>{"".join(parts)}</m:oMath>'
    paragraph._p.append(parse_xml(xml))


def add_example_derivative(doc, func_label, a_label, diff_num, simplified, result):
    p = doc.add_paragraph()
    p.add_run("例：").bold = True
    p.add_run(func_label)
    p.add_run(" について、定義から微分係数を求める。")
    add_equation(
        doc,
        [
            m_text("f'("),
            m_text(a_label),
            m_text(")="),
            m_lim_low("h→0"),
            m_frac(m_text(diff_num), m_text("h")),
        ],
        center=True,
    )
    add_equation(doc, [m_text("="), m_lim_low("h→0"), m_text(simplified), m_text("="), m_text(result)], center=True)


def make_worksheet():
    doc = setup_doc("数学Ⅲ 第03回　定義から微分係数を求めるワークシート")
    add_footer(doc, "生徒用ワークシート")

    add_note_box(
        doc,
        "このプリントのゴール",
        [
            "数学Ⅱの計算を復習しながら、最後に数Ⅲの微分係数を定義から求めます。",
            "元の課題と同じ問題は使いません。すべて類似問題で練習します。",
        ],
    )
    add_checklist(
        doc,
        [
            "分数式を正しく計算できる",
            "有理化ができる",
            "なぜこの計算をするのかを説明できる",
            "微分係数を定義から求められる",
        ],
    )

    add_section(doc, "STEP1　数Ⅱ復習①　分数式")
    doc.add_paragraph("分数は、分母が違うとそのまま引けません。分母は「1つ分の大きさ」を表しているので、まず単位をそろえます。")
    add_equation(doc, [m_frac(m_text("1"), m_text("x+5")), m_text(" − "), m_frac(m_text("1"), m_text("5"))], center=True)
    add_note_box(doc, "通分のイメージ", ["5分の1と、x+5分の1は、同じ単位ではありません。", "引き算をするために、分母を 5(x+5) にそろえます。"])
    add_write_box(
        doc,
        [
            "① まず何をしますか。",
            "② その理由は何ですか。",
            "③ 通分した後、分子はどうなりますか。",
        ],
    )
    add_subsection(doc, "練習")
    fraction_practice = [
        (m_frac(m_text("1"), m_text("x+h")), m_frac(m_text("1"), m_text("x")), "分母を x(x+h) にそろえる"),
        (m_frac(m_text("1"), m_text("a+h")), m_frac(m_text("1"), m_text("a")), "分母を a(a+h) にそろえる"),
        (m_frac(m_text("1"), m_text("2+h")), m_frac(m_text("1"), m_text("2")), "分母を 2(2+h) にそろえる"),
    ]
    for left, right, hint in fraction_practice:
        doc.add_paragraph(f"・次の式を通分しなさい。ヒント：{hint}")
        add_equation(doc, [left, m_text(" − "), right], center=True)
        add_write_box(doc, ["計算欄"], height_cm=1.8)

    doc.add_page_break()
    add_section(doc, "STEP2　数Ⅱ復習②　hを約分できる形にする")
    doc.add_paragraph("微分係数の定義では、最後に h → 0 とします。その前に h で割れる形を作ることが大切です。")
    add_equation(
        doc,
        [
            m_frac(m_text("1"), m_text("x+h")),
            m_text(" − "),
            m_frac(m_text("1"), m_text("x")),
            m_text(" = "),
            m_frac(m_text("x − (x+h)"), m_text("x(x+h)")),
            m_text(" = "),
            m_frac(m_text("−h"), m_text("x(x+h)")),
        ],
        center=True,
    )
    add_note_box(doc, "ここで見ていること", ["分子に h が出ると、あとで h で約分できます。", "微分の計算では、この h を作ることが大きな目的です。"], "F0FDF4")
    add_write_box(
        doc,
        [
            "① なぜ通分しますか。",
            "② 分子の x − (x+h) は何になりますか。",
            "③ どの h が約分できますか。",
        ],
    )
    doc.add_paragraph("練習：次の形を h で約分できるところまで変形しなさい。")
    h_practice = [
        (m_frac(m_text("1"), m_text("3+h")), m_frac(m_text("1"), m_text("3"))),
        (m_frac(m_text("1"), m_text("5+h")), m_frac(m_text("1"), m_text("5"))),
        (m_frac(m_text("1"), m_text("x+h")), m_frac(m_text("1"), m_text("x"))),
    ]
    for left, right in h_practice:
        doc.add_paragraph("・次の式を h で約分できるところまで変形しなさい。")
        add_equation(doc, [m_frac(left + m_text(" − ") + right, m_text("h"))], center=True)
        add_write_box(doc, ["計算欄"], height_cm=1.8)

    doc.add_page_break()
    add_section(doc, "STEP3　数Ⅱ復習③　有理化")
    doc.add_paragraph("平方根を含む式では、引き算したままだと h が見えにくいことがあります。そこで共役な式を掛けます。")
    add_equation(doc, [m_sqrt(m_text("9+h")), m_text(" − 3")], center=True)
    add_note_box(doc, "なぜ有理化するのか", ["平方根の引き算は、そのままでは h が取り出しにくいです。", "共役な式を掛けると、平方根が消えて h が見えるようになります。"])
    add_equation(
        doc,
        [
            m_frac(m_sqrt(m_text("9+h")) + m_text(" − 3"), m_text("h")),
            m_text(" × "),
            m_frac(m_sqrt(m_text("9+h")) + m_text(" + 3"), m_sqrt(m_text("9+h")) + m_text(" + 3")),
        ],
        center=True,
    )
    add_write_box(
        doc,
        [
            "① 掛ける相手は何ですか。",
            "② 分母と分子の両方に同じ式を掛ける理由は何ですか。",
            "③ 平方根が消えた後、分子は何になりますか。",
        ],
    )
    rationalize_practice = [
        (m_sqrt(m_text("4+h")), m_text("2")),
        (m_sqrt(m_text("16+h")), m_text("4")),
        (m_sqrt(m_text("x+h")), m_sqrt(m_text("x"))),
    ]
    for left, right in rationalize_practice:
        doc.add_paragraph("・次の式を有理化しなさい。")
        add_equation(doc, [m_frac(left + m_text(" − ") + right, m_text("h"))], center=True)
        add_write_box(doc, ["計算欄"], height_cm=1.8)

    doc.add_page_break()
    add_section(doc, "STEP4　微分係数とは？")
    doc.add_paragraph("微分係数は、ある1点での「瞬間の変化の割合」です。最初から瞬間は見えないので、近い2点の平均の変化を考えます。")
    add_equation(
        doc,
        [
            m_text("平均の変化 = "),
            m_frac(m_text("f(a+h) − f(a)"), m_text("h")),
        ],
        center=True,
    )
    add_equation(
        doc,
        [
            m_text("微分係数 = "),
            m_lim_low("h→0"),
            m_frac(m_text("f(a+h) − f(a)"), m_text("h")),
        ],
        center=True,
    )
    add_write_box(
        doc,
        [
            "① f(a+h) − f(a) は何を表していますか。",
            "② hで割る理由は何ですか。",
            "③ h→0 とする理由を言葉で説明しましょう。",
        ],
    )

    doc.add_page_break()
    add_section(doc, "STEP5　類似問題①　分数式の微分係数")
    doc.add_paragraph("問題：次の関数について、定義から微分係数を求めなさい。")
    add_equation(doc, [m_text("f(x)="), m_frac(m_text("1"), m_text("x+2")), m_text(",  f'(3)")], center=True)
    add_equation(doc, [m_text("f'(3)="), m_lim_low("h→0"), m_frac(m_text("f(3+h)−f(3)"), m_text("h"))], center=True)
    add_write_box(
        doc,
        [
            "① f(3+h) を書きましょう。",
            "② f(3) を書きましょう。",
            "③ 分母をそろえましょう。",
            "④ hで約分しましょう。",
            "⑤ h→0 を代入しましょう。",
        ],
        height_cm=1.7,
    )

    add_section(doc, "STEP6　類似問題②　平方根の微分係数")
    doc.add_paragraph("問題：次の関数について、定義から微分係数を求めなさい。")
    add_equation(doc, [m_text("f(x)="), m_sqrt(m_text("x+4")), m_text(",  f'(5)")], center=True)
    add_equation(doc, [m_text("f'(5)="), m_lim_low("h→0"), m_frac(m_sqrt(m_text("9+h")) + m_text("−3"), m_text("h"))], center=True)
    add_write_box(
        doc,
        [
            "① どの式を掛けて有理化しますか。",
            "② 分子は何になりますか。",
            "③ hで約分しましょう。",
            "④ h→0 を代入しましょう。",
        ],
        height_cm=1.7,
    )

    doc.add_page_break()
    add_section(doc, "STEP7　自分で考える問題")
    thinking_questions = [
        [m_text("f(x)="), m_frac(m_text("1"), m_text("x+4")), m_text(",  f'(1)")],
        [m_text("f(x)="), m_sqrt(m_text("x+5")), m_text(",  f'(4)")],
    ]
    for q in thinking_questions:
        doc.add_paragraph("次の関数について、定義から微分係数を求めなさい。")
        add_equation(doc, q, center=True)
        add_write_box(doc, ["方針", "計算", "答え"], height_cm=2.0)

    add_section(doc, "STEP8　本番チャレンジ")
    doc.add_paragraph("問題：次の関数について、定義から微分係数を求めなさい。")
    add_equation(doc, [m_text("f(x)="), m_frac(m_text("1"), m_text("x+1")), m_text(",  f'(2)")], center=True)
    add_write_box(
        doc,
        [
            "① 定義の式を書く",
            "② f(2+h) と f(2) を代入する",
            "③ 通分する",
            "④ hで約分する",
            "⑤ h→0 を計算する",
            "⑥ 答え",
        ],
        height_cm=1.55,
    )
    doc.add_paragraph("最後に、自分の言葉で説明しましょう。")
    add_write_box(doc, ["なぜ通分したのか", "なぜhで約分したのか", "なぜh→0にしたのか"], height_cm=1.5)
    return doc


def make_hints():
    doc = setup_doc("数学Ⅲ 第03回　ヒント集")
    add_footer(doc, "ヒント集")
    add_note_box(doc, "使い方", ["答えを見る前に、次に何をすればよいかだけを確認するための冊子です。", "計算そのものは自分で進めましょう。"])

    hints = [
        ("STEP1 分数式", ["分母が違う分数は、そのまま引けません。", "共通の分母を先に決めます。", "次の2つの分数なら、共通分母は x(x+h) です。"]),
        ("STEP2 hを作る", ["分子に h が出る形を目指します。", "x − (x+h) の括弧を外すと、x − x − h になります。", "最後に分子の h と、定義の分母にある h が約分できます。"]),
        ("STEP3 有理化", ["平方根の引き算では、共役な式を掛けます。", "差の式に対して、符号だけを変えた和の式を掛けます。", "(A−B)(A+B)=A²−B² を使います。"]),
        ("STEP4 微分係数", ["まず f(a+h) と f(a) を別々に書きます。", "差 f(a+h)−f(a) を作ります。", "hで割って、最後に h→0 とします。"]),
        ("STEP5 分数式の類似問題", ["f(3+h) と f(3) を先に書きます。", "2つの分数の差を通分します。", "分子に −h が出たら、hで約分できます。"]),
        ("STEP6 平方根の類似問題", ["f(5+h) と f(5) を先に書きます。", "平方根を含む差に、共役な式を掛けます。", "分子は (9+h)−9 になるので h だけが残ります。"]),
        ("STEP8 本番チャレンジ", ["f(2+h) と f(2) を先に書きます。", "共通分母は 3(3+h) です。", "最後は分子に h が残らない形にしてから h→0 を考えます。"]),
    ]
    for title, lines in hints:
        add_section(doc, title)
        for i, line in enumerate(lines, 1):
            doc.add_paragraph(f"ヒント{i}：{line}")
            if title == "STEP1 分数式" and i == 3:
                add_equation(doc, [m_frac(m_text("1"), m_text("x+h")), m_text(" と "), m_frac(m_text("1"), m_text("x"))], center=True)
            if title == "STEP3 有理化" and i == 2:
                add_equation(doc, [m_sqrt(m_text("9+h")), m_text(" − 3 と "), m_sqrt(m_text("9+h")), m_text(" + 3")], center=True)
            if title == "STEP5 分数式の類似問題" and i == 1:
                add_equation(doc, [m_text("f(3+h)="), m_frac(m_text("1"), m_text("5+h")), m_text(",  f(3)="), m_frac(m_text("1"), m_text("5"))], center=True)
            if title == "STEP5 分数式の類似問題" and i == 2:
                add_equation(doc, [m_frac(m_text("1"), m_text("5+h")), m_text(" − "), m_frac(m_text("1"), m_text("5"))], center=True)
            if title == "STEP6 平方根の類似問題" and i == 1:
                add_equation(doc, [m_text("f(5+h)="), m_sqrt(m_text("9+h")), m_text(",  f(5)=3")], center=True)
            if title == "STEP6 平方根の類似問題" and i == 2:
                add_equation(doc, [m_sqrt(m_text("9+h")), m_text("−3 に "), m_sqrt(m_text("9+h")), m_text("+3 を掛ける")], center=True)
            if title == "STEP8 本番チャレンジ" and i == 1:
                add_equation(doc, [m_text("f(2+h)="), m_frac(m_text("1"), m_text("3+h")), m_text(",  f(2)="), m_frac(m_text("1"), m_text("3"))], center=True)
        doc.add_paragraph()
    return doc


def make_answers():
    doc = setup_doc("数学Ⅲ 第03回　解答・解説集")
    add_footer(doc, "解答・解説集")
    add_note_box(doc, "この冊子の方針", ["途中式を省略せず、なぜその式変形をするのかを説明します。"])

    add_section(doc, "STEP1・STEP2　分数式")
    add_equation(
        doc,
        [
            m_frac(m_text("1"), m_text("x+h")),
            m_text(" − "),
            m_frac(m_text("1"), m_text("x")),
            m_text(" = "),
            m_frac(m_text("x−(x+h)"), m_text("x(x+h)")),
            m_text(" = "),
            m_frac(m_text("−h"), m_text("x(x+h)")),
        ],
        center=True,
    )
    doc.add_paragraph("分母が違うので通分します。分子は x−(x+h)=−h です。この h が、微分係数の定義にある分母の h と約分されます。")

    add_section(doc, "STEP3　有理化")
    add_equation(
        doc,
        [
            m_frac(m_sqrt(m_text("9+h")) + m_text("−3"), m_text("h")),
            m_text("×"),
            m_frac(m_sqrt(m_text("9+h")) + m_text("+3"), m_sqrt(m_text("9+h")) + m_text("+3")),
            m_text("="),
            m_frac(m_text("h"), m_text("h") + m_text("(") + m_sqrt(m_text("9+h")) + m_text("+3)")),
        ],
        center=True,
    )
    doc.add_paragraph("共役な式を掛けると、分子は (9+h)−9=h になります。これで h が約分できます。")

    add_section(doc, "STEP5　分数式の微分係数")
    add_equation(doc, [m_text("f(x)="), m_frac(m_text("1"), m_text("x+2")), m_text(",  f'(3)")], center=True)
    add_equation(doc, [m_text("f(3+h)="), m_frac(m_text("1"), m_text("5+h")), m_text(",  f(3)="), m_frac(m_text("1"), m_text("5"))], center=True)
    add_equation(
        doc,
        [
            m_text("f'(3)="),
            m_lim_low("h→0"),
            m_frac(m_frac(m_text("1"), m_text("5+h")) + m_text("−") + m_frac(m_text("1"), m_text("5")), m_text("h")),
            m_text("="),
            m_lim_low("h→0"),
            m_frac(m_text("−1"), m_text("5(5+h)")),
            m_text("=−"),
            m_frac(m_text("1"), m_text("25")),
        ],
        center=True,
    )
    doc.add_paragraph("通分すると分子は 5−(5+h)=−h です。hで約分してから h→0 を代入します。")

    add_section(doc, "STEP6　平方根を含む微分係数")
    add_equation(doc, [m_text("f(x)="), m_sqrt(m_text("x+4")), m_text(",  f'(5)")], center=True)
    add_equation(doc, [m_text("f(5+h)="), m_sqrt(m_text("9+h")), m_text(",  f(5)=3")], center=True)
    add_equation(
        doc,
        [
            m_text("f'(5)="),
            m_lim_low("h→0"),
            m_frac(m_sqrt(m_text("9+h")) + m_text("−3"), m_text("h")),
            m_text("="),
            m_lim_low("h→0"),
            m_frac(m_text("1"), m_sqrt(m_text("9+h")) + m_text("+3")),
            m_text("="),
            m_frac(m_text("1"), m_text("6")),
        ],
        center=True,
    )
    doc.add_paragraph("有理化で分子を h にし、hで約分します。最後に h=0 を入れると分母は 3+3=6 です。")

    add_section(doc, "STEP7　自分で考える問題")
    doc.add_paragraph("1. 次の関数の微分係数")
    add_equation(doc, [m_text("f(x)="), m_frac(m_text("1"), m_text("x+4")), m_text(",  f'(1)")], center=True)
    add_equation(doc, [m_text("答え：−"), m_frac(m_text("1"), m_text("25"))], center=True)
    doc.add_paragraph("2. 次の関数の微分係数")
    add_equation(doc, [m_text("f(x)="), m_sqrt(m_text("x+5")), m_text(",  f'(4)")], center=True)
    add_equation(doc, [m_text("答え："), m_frac(m_text("1"), m_text("6"))], center=True)

    add_section(doc, "STEP8　本番チャレンジ")
    doc.add_paragraph("次の関数の微分係数")
    add_equation(doc, [m_text("f(x)="), m_frac(m_text("1"), m_text("x+1")), m_text(",  f'(2)")], center=True)
    add_equation(
        doc,
        [
            m_text("f'(2)="),
            m_lim_low("h→0"),
            m_frac(m_frac(m_text("1"), m_text("3+h")) + m_text("−") + m_frac(m_text("1"), m_text("3")), m_text("h")),
            m_text("="),
            m_lim_low("h→0"),
            m_frac(m_text("−1"), m_text("3(3+h)")),
            m_text("=−"),
            m_frac(m_text("1"), m_text("9")),
        ],
        center=True,
    )
    doc.add_paragraph("答えは次の通りです。")
    add_equation(doc, [m_text("−"), m_frac(m_text("1"), m_text("9"))], center=True)
    return doc


def make_teacher():
    doc = setup_doc("数学Ⅲ 第03回　教師用指導ポイント")
    add_footer(doc, "教師用指導ポイント")
    add_note_box(
        doc,
        "授業設計のねらい",
        [
            "数学Ⅱの式変形を復習しながら、数Ⅲの微分係数の定義へ接続します。",
            "元問題の穴埋めを直接練習するのではなく、類似問題で計算の理由を言語化させます。",
        ],
    )
    add_section(doc, "想定時間")
    doc.add_paragraph("自習：約2時間。授業内で扱う場合は、STEP1〜3を30分、STEP4〜6を35分、STEP7〜8を25分、振り返りを10分程度。")

    add_section(doc, "つまずきやすい箇所")
    items = [
        ("通分", "分母をそろえる目的がわからず、分子だけを引いてしまう。"),
        ("括弧外し", "x−(x+h) を x−x+h としてしまう。"),
        ("hの約分", "h→0 を先に代入して未定形になり、計算が止まる。"),
        ("有理化", "共役を分子だけに掛けてしまう。"),
        ("微分係数の意味", "公式暗記になり、平均の変化から瞬間へ近づける流れを説明できない。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "項目"
    table.rows[0].cells[1].text = "よくあるつまずき"
    table.rows[0].cells[2].text = "声かけ"
    for k, v in items:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v
        row[2].text = "何のためにその計算をするのかを、先に言葉で確認させる。"
    doc.add_paragraph()

    add_section(doc, "よくある誤答と対応")
    for title, body in [
        ("f(a+h) の代入ミス", "a に h を足すだけでなく、関数全体の x に a+h を代入することを確認する。"),
        ("hで割る位置のミス", "分子全体 f(a+h)−f(a) を h で割る。差を作る前に一部だけ割らない。"),
        ("有理化後の符号ミス", "(A−B)(A+B)=A²−B² を書かせ、どちらが A かを確認する。"),
    ]:
        add_subsection(doc, title)
        doc.add_paragraph(body)

    add_section(doc, "数学Ⅱとの関連")
    doc.add_paragraph("分数式の計算、通分、約分、因数分解、有理化はすべて数学Ⅱまでの計算です。数Ⅲの新しさは、これらを h→0 の極限に接続する点にあります。")
    add_section(doc, "教科書を読み返す観点")
    doc.add_paragraph("微分係数の定義、導関数の定義、分数関数・平方根を含む関数の導関数の導入部分を確認させます。生徒には『公式の前に、定義の式がどう変形されるか』を見るよう指示します。")
    add_section(doc, "評価の観点")
    add_checklist(
        doc,
        [
            "通分の理由を説明できている",
            "hで約分する前に h→0 を代入していない",
            "有理化で共役を選べている",
            "微分係数を平均の変化から瞬間の変化へ近づけるものとして説明できている",
        ],
    )
    return doc


def main():
    OUT_DIR.mkdir(exist_ok=True)
    docs = [
        ("01_生徒用ワークシート.docx", make_worksheet()),
        ("02_ヒント集.docx", make_hints()),
        ("03_解答解説集.docx", make_answers()),
        ("04_教師用指導ポイント.docx", make_teacher()),
    ]
    for filename, doc in docs:
        doc.save(OUT_DIR / filename)
        print(OUT_DIR / filename)


if __name__ == "__main__":
    main()
